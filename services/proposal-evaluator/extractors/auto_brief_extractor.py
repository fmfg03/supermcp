"""
Auto-Brief Extractor - Zero-Touch Onboarding
==========================================

Extracts project context from PDFs, links, voice notes, emails, and documents.
Automatically generates project manifests and launches evaluation cycles.
"""

import asyncio
import re
import json
from typing import Dict, Any, List, Optional, Union, Tuple
from dataclasses import dataclass, asdict
from datetime import datetime
from pathlib import Path
import hashlib
import tempfile
import aiofiles
import aiohttp

# Document processing
import PyPDF2
from docx import Document
import speech_recognition as sr
from pydub import AudioSegment

# Web scraping
from bs4 import BeautifulSoup
import requests

# Email processing
import email
from email.mime.text import MIMEText
import imaplib

from ..utils.llm_client import LLMClient
from ..utils.logger import setup_logger
from ..main import ProposalEvaluator, ProposalContext

logger = setup_logger(__name__)

@dataclass
class ExtractedBrief:
    """Extracted project brief information"""
    source_type: str  # "pdf", "url", "voice", "email", "document"
    source_location: str
    client_name: str
    project_description: str
    business_context: str
    objectives: List[str]
    constraints: List[str]
    budget_info: Optional[str]
    timeline_info: Optional[str]
    stakeholders: List[str]
    success_metrics: List[str]
    technical_requirements: List[str]
    compliance_requirements: List[str]
    extracted_at: datetime
    confidence_score: float
    raw_content: str

@dataclass
class ProcessingResult:
    """Result of auto-brief processing"""
    brief: ExtractedBrief
    proposal_context: ProposalContext
    auto_evaluation_started: bool
    evaluation_id: Optional[str]
    processing_log: List[str]

class DocumentProcessor:
    """Handles document processing (PDF, DOCX, TXT)"""
    
    def __init__(self, llm_client: LLMClient):
        self.llm_client = llm_client
    
    async def process_pdf(self, file_path: str) -> str:
        """Extract text from PDF"""
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                
                logger.info(f"Extracted {len(text)} characters from PDF")
                return text.strip()
                
        except Exception as e:
            logger.error(f"PDF processing failed: {e}")
            return ""
    
    async def process_docx(self, file_path: str) -> str:
        """Extract text from DOCX"""
        
        try:
            doc = Document(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            logger.info(f"Extracted {len(text)} characters from DOCX")
            return text.strip()
            
        except Exception as e:
            logger.error(f"DOCX processing failed: {e}")
            return ""
    
    async def process_text_file(self, file_path: str) -> str:
        """Process plain text file"""
        
        try:
            async with aiofiles.open(file_path, mode='r', encoding='utf-8') as f:
                content = await f.read()
                
            logger.info(f"Extracted {len(content)} characters from text file")
            return content.strip()
            
        except Exception as e:
            logger.error(f"Text file processing failed: {e}")
            return ""

class VoiceProcessor:
    """Handles voice/audio processing"""
    
    def __init__(self, llm_client: LLMClient):
        self.llm_client = llm_client
        self.recognizer = sr.Recognizer()
    
    async def process_audio(self, file_path: str) -> str:
        """Extract text from audio file"""
        
        try:
            # Convert audio to WAV if needed
            wav_path = await self._convert_to_wav(file_path)
            
            # Transcribe audio
            with sr.AudioFile(wav_path) as source:
                audio = self.recognizer.record(source)
            
            # Use Google Speech Recognition
            text = self.recognizer.recognize_google(audio)
            
            logger.info(f"Transcribed {len(text)} characters from audio")
            return text
            
        except Exception as e:
            logger.error(f"Audio processing failed: {e}")
            return ""
    
    async def _convert_to_wav(self, file_path: str) -> str:
        """Convert audio file to WAV format"""
        
        file_ext = Path(file_path).suffix.lower()
        
        if file_ext == '.wav':
            return file_path
        
        try:
            # Create temporary WAV file
            temp_wav = tempfile.mktemp(suffix='.wav')
            
            # Convert using pydub
            if file_ext in ['.mp3', '.m4a', '.ogg']:
                audio = AudioSegment.from_file(file_path)
                audio.export(temp_wav, format="wav")
                return temp_wav
            else:
                logger.warning(f"Unsupported audio format: {file_ext}")
                return file_path
                
        except Exception as e:
            logger.error(f"Audio conversion failed: {e}")
            return file_path

class WebExtractor:
    """Handles web content extraction"""
    
    def __init__(self, llm_client: LLMClient):
        self.llm_client = llm_client
    
    async def extract_from_url(self, url: str) -> str:
        """Extract content from web URL"""
        
        try:
            async with aiohttp.ClientSession() as session:
                async with session.get(url, timeout=30) as response:
                    if response.status == 200:
                        html_content = await response.text()
                        
                        # Parse with BeautifulSoup
                        soup = BeautifulSoup(html_content, 'html.parser')
                        
                        # Remove script and style elements
                        for script in soup(["script", "style"]):
                            script.decompose()
                        
                        # Extract text
                        text = soup.get_text()
                        
                        # Clean up whitespace
                        lines = (line.strip() for line in text.splitlines())
                        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
                        text = ' '.join(chunk for chunk in chunks if chunk)
                        
                        logger.info(f"Extracted {len(text)} characters from URL")
                        return text
                    else:
                        logger.error(f"Failed to fetch URL: {response.status}")
                        return ""
                        
        except Exception as e:
            logger.error(f"URL extraction failed: {e}")
            return ""

class EmailProcessor:
    """Handles email processing"""
    
    def __init__(self, llm_client: LLMClient):
        self.llm_client = llm_client
    
    async def process_email(self, email_content: str) -> str:
        """Process email content"""
        
        try:
            # Parse email
            msg = email.message_from_string(email_content)
            
            # Extract text content
            text_content = ""
            
            if msg.is_multipart():
                for part in msg.walk():
                    if part.get_content_type() == "text/plain":
                        text_content += part.get_payload(decode=True).decode('utf-8', errors='ignore')
            else:
                if msg.get_content_type() == "text/plain":
                    text_content = msg.get_payload(decode=True).decode('utf-8', errors='ignore')
            
            # Extract headers for context
            subject = msg.get('Subject', '')
            sender = msg.get('From', '')
            date = msg.get('Date', '')
            
            # Combine into structured text
            full_content = f"""
Email Subject: {subject}
From: {sender}
Date: {date}

Content:
{text_content}
"""
            
            logger.info(f"Processed email with {len(full_content)} characters")
            return full_content.strip()
            
        except Exception as e:
            logger.error(f"Email processing failed: {e}")
            return email_content  # Fallback to raw content

class AutoBriefExtractor:
    """
    Main extractor that processes multiple input types and generates project briefs
    """
    
    def __init__(self, config: Dict[str, Any]):
        self.config = config
        self.llm_client = LLMClient(config.get("llm", {}))
        
        # Initialize processors
        self.document_processor = DocumentProcessor(self.llm_client)
        self.voice_processor = VoiceProcessor(self.llm_client)
        self.web_extractor = WebExtractor(self.llm_client)
        self.email_processor = EmailProcessor(self.llm_client)
        
        # Initialize evaluator for auto-launch
        self.evaluator = ProposalEvaluator()
    
    async def extract_from_source(
        self, 
        source: Union[str, Path], 
        source_type: Optional[str] = None,
        auto_evaluate: bool = True
    ) -> ProcessingResult:
        """
        Extract project brief from any source type
        
        Args:
            source: File path, URL, or content string
            source_type: Optional type hint ("pdf", "url", "voice", "email", "auto")
            auto_evaluate: Whether to automatically start evaluation
        """
        
        processing_log = []
        processing_log.append(f"Started processing at {datetime.utcnow()}")
        
        try:
            # Determine source type if not provided
            if source_type is None:
                source_type = await self._detect_source_type(source)
            
            processing_log.append(f"Detected source type: {source_type}")
            
            # Extract raw content
            raw_content = await self._extract_raw_content(source, source_type)
            
            if not raw_content:
                raise ValueError("No content could be extracted from source")
            
            processing_log.append(f"Extracted {len(raw_content)} characters of raw content")
            
            # Extract structured brief
            brief = await self._extract_structured_brief(raw_content, source_type, str(source))
            
            processing_log.append(f"Generated structured brief with confidence {brief.confidence_score}")
            
            # Create proposal context
            proposal_context = await self._create_proposal_context(brief)
            
            processing_log.append("Created proposal context")
            
            # Auto-evaluate if requested
            evaluation_id = None
            auto_evaluation_started = False
            
            if auto_evaluate and brief.confidence_score > 0.6:
                try:
                    evaluation_id = await self._launch_auto_evaluation(proposal_context)
                    auto_evaluation_started = True
                    processing_log.append(f"Launched auto-evaluation: {evaluation_id}")
                except Exception as e:
                    processing_log.append(f"Auto-evaluation failed: {e}")
            
            return ProcessingResult(
                brief=brief,
                proposal_context=proposal_context,
                auto_evaluation_started=auto_evaluation_started,
                evaluation_id=evaluation_id,
                processing_log=processing_log
            )
            
        except Exception as e:
            processing_log.append(f"Processing failed: {e}")
            logger.error(f"Auto-brief extraction failed: {e}")
            raise
    
    async def _detect_source_type(self, source: Union[str, Path]) -> str:
        """Auto-detect source type"""
        
        source_str = str(source)
        
        # Check if it's a URL
        if source_str.startswith(('http://', 'https://')):
            return "url"
        
        # Check if it's a file path
        if Path(source).exists():
            file_path = Path(source)
            ext = file_path.suffix.lower()
            
            if ext == '.pdf':
                return "pdf"
            elif ext in ['.docx', '.doc']:
                return "docx"
            elif ext in ['.txt', '.md']:
                return "text"
            elif ext in ['.wav', '.mp3', '.m4a', '.ogg']:
                return "voice"
            elif ext in ['.eml', '.msg']:
                return "email"
        
        # Check if it looks like email content
        if 'Subject:' in source_str and 'From:' in source_str:
            return "email"
        
        # Default to text content
        return "text"
    
    async def _extract_raw_content(self, source: Union[str, Path], source_type: str) -> str:
        """Extract raw content based on source type"""
        
        if source_type == "pdf":
            return await self.document_processor.process_pdf(str(source))
        
        elif source_type == "docx":
            return await self.document_processor.process_docx(str(source))
        
        elif source_type == "text":
            if Path(source).exists():
                return await self.document_processor.process_text_file(str(source))
            else:
                return str(source)  # Treat as raw text
        
        elif source_type == "voice":
            return await self.voice_processor.process_audio(str(source))
        
        elif source_type == "url":
            return await self.web_extractor.extract_from_url(str(source))
        
        elif source_type == "email":
            if Path(source).exists():
                with open(source, 'r') as f:
                    email_content = f.read()
            else:
                email_content = str(source)
            
            return await self.email_processor.process_email(email_content)
        
        else:
            return str(source)  # Fallback to raw content
    
    async def _extract_structured_brief(
        self, 
        raw_content: str, 
        source_type: str, 
        source_location: str
    ) -> ExtractedBrief:
        """Extract structured project brief from raw content"""
        
        # Build extraction prompt
        extraction_prompt = f"""Extract a structured project brief from the following content:

SOURCE TYPE: {source_type}
CONTENT:
{raw_content}

Extract the following information in JSON format:
{{
    "client_name": "extracted or inferred client/company name",
    "project_description": "clear description of the project",
    "business_context": "business background and situation",
    "objectives": ["business objective 1", "business objective 2"],
    "constraints": ["constraint 1", "constraint 2"],
    "budget_info": "budget information if mentioned",
    "timeline_info": "timeline information if mentioned", 
    "stakeholders": ["stakeholder 1", "stakeholder 2"],
    "success_metrics": ["metric 1", "metric 2"],
    "technical_requirements": ["tech requirement 1", "tech requirement 2"],
    "compliance_requirements": ["compliance req 1", "compliance req 2"],
    "confidence_score": 0.8
}}

Rules:
- Extract actual information, don't invent details
- Use "Unknown" for missing required fields
- Set confidence_score based on how complete the information is (0.0 to 1.0)
- If content is unclear, infer reasonable values but lower confidence
- Focus on business value and decision-making context"""
        
        system_prompt = """You are an expert business analyst who extracts structured project information from various content types. You understand business contexts, technical requirements, and stakeholder dynamics. Extract accurate information and assess confidence realistically."""
        
        try:
            response = await self.llm_client.generate(
                system_prompt=system_prompt,
                user_prompt=extraction_prompt,
                temperature=0.1,  # Low temperature for consistent extraction
                max_tokens=2000
            )
            
            # Parse JSON response
            try:
                extracted_data = json.loads(response)
            except json.JSONDecodeError:
                # Try to extract JSON from response if LLM added extra text
                json_match = re.search(r'\{.*\}', response, re.DOTALL)
                if json_match:
                    extracted_data = json.loads(json_match.group())
                else:
                    raise ValueError("Could not parse JSON from LLM response")
            
            # Create ExtractedBrief
            brief = ExtractedBrief(
                source_type=source_type,
                source_location=source_location,
                client_name=extracted_data.get("client_name", "Unknown"),
                project_description=extracted_data.get("project_description", ""),
                business_context=extracted_data.get("business_context", ""),
                objectives=extracted_data.get("objectives", []),
                constraints=extracted_data.get("constraints", []),
                budget_info=extracted_data.get("budget_info"),
                timeline_info=extracted_data.get("timeline_info"),
                stakeholders=extracted_data.get("stakeholders", []),
                success_metrics=extracted_data.get("success_metrics", []),
                technical_requirements=extracted_data.get("technical_requirements", []),
                compliance_requirements=extracted_data.get("compliance_requirements", []),
                extracted_at=datetime.utcnow(),
                confidence_score=float(extracted_data.get("confidence_score", 0.5)),
                raw_content=raw_content[:5000]  # Truncate for storage
            )
            
            return brief
            
        except Exception as e:
            logger.error(f"Structured extraction failed: {e}")
            
            # Fallback to basic extraction
            return ExtractedBrief(
                source_type=source_type,
                source_location=source_location,
                client_name="Unknown",
                project_description=raw_content[:500] + "..." if len(raw_content) > 500 else raw_content,
                business_context="Extracted from " + source_type,
                objectives=["Improve business operations"],
                constraints=[],
                budget_info=None,
                timeline_info=None,
                stakeholders=[],
                success_metrics=[],
                technical_requirements=[],
                compliance_requirements=[],
                extracted_at=datetime.utcnow(),
                confidence_score=0.3,  # Low confidence for fallback
                raw_content=raw_content[:5000]
            )
    
    async def _create_proposal_context(self, brief: ExtractedBrief) -> ProposalContext:
        """Create ProposalContext from extracted brief"""
        
        return ProposalContext(
            client=brief.client_name,
            context=f"{brief.business_context}\n\nProject: {brief.project_description}",
            objectives=brief.objectives if brief.objectives else ["Improve business efficiency"],
            constraints=brief.constraints,
            tone="professional"
        )
    
    async def _launch_auto_evaluation(self, context: ProposalContext) -> str:
        """Launch automatic proposal evaluation"""
        
        try:
            # Generate evaluation ID
            evaluation_id = f"auto_{datetime.utcnow().strftime('%Y%m%d_%H%M%S')}_{context.client[:10]}"
            
            # Start evaluation (this would integrate with your main evaluation system)
            result = await self.evaluator.evaluate_proposal(context)
            
            logger.info(f"Auto-evaluation launched: {evaluation_id}")
            return evaluation_id
            
        except Exception as e:
            logger.error(f"Auto-evaluation launch failed: {e}")
            raise
    
    async def batch_process_directory(self, directory_path: str) -> List[ProcessingResult]:
        """Process all files in a directory"""
        
        directory = Path(directory_path)
        results = []
        
        supported_extensions = {'.pdf', '.docx', '.doc', '.txt', '.md', '.wav', '.mp3', '.m4a', '.eml'}
        
        for file_path in directory.rglob('*'):
            if file_path.is_file() and file_path.suffix.lower() in supported_extensions:
                try:
                    logger.info(f"Processing: {file_path}")
                    result = await self.extract_from_source(file_path, auto_evaluate=False)
                    results.append(result)
                except Exception as e:
                    logger.error(f"Failed to process {file_path}: {e}")
        
        return results
    
    async def generate_summary_report(self, results: List[ProcessingResult]) -> Dict[str, Any]:
        """Generate summary report from batch processing results"""
        
        if not results:
            return {"error": "No results to summarize"}
        
        # Aggregate statistics
        total_processed = len(results)
        successful_extractions = len([r for r in results if r.brief.confidence_score > 0.5])
        auto_evaluations_started = len([r for r in results if r.auto_evaluation_started])
        
        # Source type distribution
        source_types = {}
        for result in results:
            source_type = result.brief.source_type
            source_types[source_type] = source_types.get(source_type, 0) + 1
        
        # Average confidence
        avg_confidence = sum(r.brief.confidence_score for r in results) / len(results)
        
        # Top clients
        client_counts = {}
        for result in results:
            client = result.brief.client_name
            if client != "Unknown":
                client_counts[client] = client_counts.get(client, 0) + 1
        
        top_clients = sorted(client_counts.items(), key=lambda x: x[1], reverse=True)[:5]
        
        # Common objectives
        all_objectives = []
        for result in results:
            all_objectives.extend(result.brief.objectives)
        
        from collections import Counter
        common_objectives = Counter(all_objectives).most_common(5)
        
        return {
            "summary": {
                "total_processed": total_processed,
                "successful_extractions": successful_extractions,
                "auto_evaluations_started": auto_evaluations_started,
                "average_confidence": avg_confidence,
                "success_rate": successful_extractions / total_processed
            },
            "source_distribution": source_types,
            "top_clients": dict(top_clients),
            "common_objectives": dict(common_objectives),
            "quality_metrics": {
                "high_confidence": len([r for r in results if r.brief.confidence_score > 0.8]),
                "medium_confidence": len([r for r in results if 0.5 <= r.brief.confidence_score <= 0.8]),
                "low_confidence": len([r for r in results if r.brief.confidence_score < 0.5])
            },
            "processing_timestamp": datetime.utcnow().isoformat()
        }

# Factory function
def create_auto_brief_extractor(config: Dict[str, Any]) -> AutoBriefExtractor:
    """Create an auto-brief extractor instance"""
    return AutoBriefExtractor(config)

# Convenience functions for different input types
async def extract_from_pdf(pdf_path: str, config: Dict[str, Any]) -> ProcessingResult:
    """Extract brief from PDF file"""
    extractor = create_auto_brief_extractor(config)
    return await extractor.extract_from_source(pdf_path, "pdf")

async def extract_from_url(url: str, config: Dict[str, Any]) -> ProcessingResult:
    """Extract brief from web URL"""
    extractor = create_auto_brief_extractor(config)
    return await extractor.extract_from_source(url, "url")

async def extract_from_voice(audio_path: str, config: Dict[str, Any]) -> ProcessingResult:
    """Extract brief from voice recording"""
    extractor = create_auto_brief_extractor(config)
    return await extractor.extract_from_source(audio_path, "voice")

async def extract_from_email(email_content: str, config: Dict[str, Any]) -> ProcessingResult:
    """Extract brief from email content"""
    extractor = create_auto_brief_extractor(config)
    return await extractor.extract_from_source(email_content, "email")

# CLI interface
async def extract_cli(source: str, source_type: str = None, auto_evaluate: bool = True):
    """CLI interface for brief extraction"""
    
    config = {
        "llm": {
            "provider": "openai",
            "model": "gpt-4",
            "api_key": "your_openai_key"
        }
    }
    
    extractor = create_auto_brief_extractor(config)
    
    try:
        result = await extractor.extract_from_source(
            source=source,
            source_type=source_type,
            auto_evaluate=auto_evaluate
        )
        
        print("✅ Brief extraction successful!")
        print(f"Client: {result.brief.client_name}")
        print(f"Project: {result.brief.project_description}")
        print(f"Confidence: {result.brief.confidence_score:.2f}")
        
        if result.auto_evaluation_started:
            print(f"🚀 Auto-evaluation started: {result.evaluation_id}")
        
        return result
        
    except Exception as e:
        print(f"❌ Extraction failed: {e}")
        raise